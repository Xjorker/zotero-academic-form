"""
Knowledge Base API 路由

端点：
- POST /kb/init              — 初始化知识库（建表）
- GET  /kb/schema             — 获取知识库 Schema
- POST /kb/query              — 执行 SQL 查询
- POST /kb/import/author      — 导入作者信息
- POST /kb/import/projects    — 导入项目列表
- POST /kb/import/patents     — 导入专利列表
- POST /kb/import/datasets    — 导入数据集列表
- POST /kb/import/zotero      — 从 Zotero 同步论文数据
- POST /kb/import/document    — 导入文档（分块+向量化）
- POST /kb/import/parse_cv    — 用 LLM 从 CV 文本提取结构化数据并导入知识库
- GET  /kb/documents          — 列出已导入的文档
- DELETE /kb/documents/{source} — 删除指定文档
- GET  /kb/stats              — 获取知识库统计
- POST /kb/search             — 混合检索
"""

from __future__ import annotations

import base64
import json
import os
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Request, HTTPException

from kb import (
    init_kb,
    get_kb_schema,
    execute_kb_query,
    import_author_info,
    import_projects,
    import_patents,
    import_datasets,
    import_zotero_papers,
    index_papers_to_vector_db,
)
from kb.vector_store import (
    index_document,
    index_paper_fields,
    search_paper_fields,
    delete_document,
    delete_paper_from_vector,
    hybrid_search,
    is_chroma_available,
    get_vector_stats,
    get_paper_vector_stats,
    PAPER_FIELD_TYPES,
)

router = APIRouter(prefix="/kb", tags=["knowledge-base"])


# ─── 初始化 ──────────────────────────────────────────────────

@router.post("/init")
async def kb_init():
    """初始化知识库（创建所有表）"""
    try:
        kb = init_kb()
        return {
            "success": True,
            "message": "知识库初始化完成",
            "schema": get_kb_schema(),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── Schema ──────────────────────────────────────────────────

@router.get("/schema")
async def kb_schema(db: str = "all"):
    """获取知识库 Schema"""
    try:
        schemas = get_kb_schema(db)
        return {"success": True, "tables": schemas}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── SQL 查询 ───────────────────────────────────────────────

@router.post("/query")
async def kb_query(request: Request):
    """
    执行 SQL 查询。

    请求体：
    {
        "sql": "SELECT * FROM academic_form_person",
        "db": "personal"  // "personal" / "zotero" / "auto"
    }
    """
    try:
        body = await request.json()
        sql = body.get("sql", "").strip()
        db = body.get("db", "auto")

        if not sql:
            return {"success": False, "error": "缺少 SQL 参数"}

        if not sql.upper().strip().startswith("SELECT"):
            return {"success": False, "error": "只允许 SELECT 查询"}

        results = execute_kb_query(sql, db=db)

        return {
            "success": True,
            "rows": results,
            "total": len(results),
            "db": db,
        }
    except ValueError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── 数据导入 ───────────────────────────────────────────────

@router.post("/import/author")
async def kb_import_author(request: Request):
    """
    导入作者信息（覆盖写入）。

    请求体：{ "orcid": "", "name": "", "affiliation": "", ... }
    """
    try:
        body = await request.json()
        count = import_author_info(body)
        return {"success": True, "message": f"作者信息已导入", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/import/projects")
async def kb_import_projects(request: Request):
    """
    导入项目列表（覆盖写入）。

    请求体：{ "projects": [{ "title": "", "year": "", "funder": "", ... }] }
    """
    try:
        body = await request.json()
        projects = body.get("projects", [])
        if not isinstance(projects, list):
            return {"success": False, "error": "projects 必须是数组"}
        count = import_projects(projects)
        return {"success": True, "message": f"已导入 {count} 个项目", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/import/patents")
async def kb_import_patents(request: Request):
    """
    导入专利列表（覆盖写入）。

    请求体：{ "patents": [{ "title": "", "year": "", "office": "", ... }] }
    """
    try:
        body = await request.json()
        patents = body.get("patents", [])
        if not isinstance(patents, list):
            return {"success": False, "error": "patents 必须是数组"}
        count = import_patents(patents)
        return {"success": True, "message": f"已导入 {count} 条专利", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/import/datasets")
async def kb_import_datasets(request: Request):
    """
    导入数据集列表（覆盖写入）。

    请求体：{ "datasets": [{ "title": "", "year": "", "repo": "", ... }] }
    """
    try:
        body = await request.json()
        datasets = body.get("datasets", [])
        if not isinstance(datasets, list):
            return {"success": False, "error": "datasets 必须是数组"}
        count = import_datasets(datasets)
        return {"success": True, "message": f"已导入 {count} 条数据集", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/import/zotero")
async def kb_import_zotero(request: Request):
    """
    从 Zotero 同步论文数据（增量更新）。

    请求体：
    {
        "items": [
            {
                "key": "ABCD1234",
                "itemType": "journalArticle",
                "title": "...",
                "DOI": "...",
                "year": 2024,
                "creators": [{"creatorType": "author", "firstName": "", "lastName": ""}],
                "tags": [{"tag": "knowledge graph"}]
            }
        ]
    }
    """
    try:
        body = await request.json()
        items = body.get("items", [])
        if not isinstance(items, list):
            return {"success": False, "error": "items 必须是数组"}

        stats = import_zotero_papers(items)
        return {
            "success": True,
            "message": f"同步完成",
            "stats": stats,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/import/document")
async def kb_import_document(request: Request):
    """
    导入文档到知识库（自动分块 + 向量化）。

    支持两种模式：
    1. PDF/DOCX 文件：传 file_base64（Base64 编码的二进制文件），后端自动提取文本
    2. 纯文本：传 content（直接文本内容）

    请求体（模式1 — 文件上传）：
    {
        "source": "cv_zhangsan.pdf",
        "file_base64": "JVBERi0xLjQK...",  // Base64 编码的文件二进制
        "doc_type": "cv",
        "title": "张三的个人简历",
        "chunk_size": 500,
        "chunk_overlap": 100,
        "metadata": {}
    }

    请求体（模式2 — 纯文本）：
    {
        "source": "notes.txt",
        "content": "大段文本内容...",
        "doc_type": "general",
        "title": "笔记",
        "chunk_size": 500,
        "chunk_overlap": 100,
        "metadata": {}
    }
    """
    try:
        body = await request.json()
        source = body.get("source", "")
        file_base64 = body.get("file_base64", "")
        content = body.get("content", "")
        doc_type = body.get("doc_type", "general")
        title = body.get("title", source)
        chunk_size = body.get("chunk_size", 500)
        chunk_overlap = body.get("chunk_overlap", 100)
        metadata = body.get("metadata", {})

        if not source:
            return {"success": False, "error": "缺少 source"}

        # 模式1: 文件上传（Base64），需要后端提取文本
        if file_base64:
            ext = os.path.splitext(source)[1].lower()
            content = _extract_text_from_file(file_base64, ext)
            if not content or not content.strip():
                return {"success": False, "error": f"无法从 {ext} 文件中提取文本"}
            print(f"[KB] 文件文本提取完成: {len(content)} 字符 (来自 {source})")

        if not content or not content.strip():
            return {"success": False, "error": "缺少 content 或文件提取结果为空"}

        # 文本分块
        chunks = _chunk_text(content.strip(), chunk_size, chunk_overlap)
        if not chunks:
            return {"success": False, "error": "分块结果为空"}

        # 索引到 SQLite + Chroma
        result = index_document(
            source=source,
            chunks=chunks,
            doc_type=doc_type,
            title=title,
            metadata=metadata,
            replace=True,
        )

        return {
            "success": True,
            "message": f"文档已导入: {len(chunks)} 个分块",
            "chunks": len(chunks),
            "extracted_chars": len(content),
            "sqlite_chunks": result["sqlite_chunks"],
            "chroma_chunks": result["chroma_chunks"],
            "chroma_available": result["chroma_available"],
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


@router.post("/import/chunks")
async def kb_import_chunks(request: Request):
    """
    直接导入预分块的文本（跳过分块步骤）。

    请求体：
    {
        "source": "manual_entry_1",
        "chunks": ["第一段内容...", "第二段内容..."],
        "doc_type": "general",
        "title": "手动录入",
        "metadata": {}
    }
    """
    try:
        body = await request.json()
        source = body.get("source", "")
        chunks = body.get("chunks", [])
        doc_type = body.get("doc_type", "general")
        title = body.get("title", "")
        metadata = body.get("metadata", {})

        if not source or not chunks:
            return {"success": False, "error": "缺少 source 或 chunks"}

        result = index_document(
            source=source,
            chunks=chunks,
            doc_type=doc_type,
            title=title,
            metadata=metadata,
            replace=True,
        )

        return {
            "success": True,
            "message": f"已导入 {len(chunks)} 个分块",
            "chunks": len(chunks),
            "sqlite_chunks": result["sqlite_chunks"],
            "chroma_chunks": result["chroma_chunks"],
            "chroma_available": result["chroma_available"],
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── 文档管理 ───────────────────────────────────────────────

@router.get("/documents")
async def kb_list_documents():
    """列出所有已导入的文档来源"""
    try:
        from kb.database import get_kb
        kb = get_kb()
        sources = kb.list_document_sources()
        return {"success": True, "documents": sources}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.delete("/documents/{source:path}")
async def kb_delete_document(source: str):
    """删除指定来源的所有文档（SQLite + Chroma）"""
    try:
        result = delete_document(source)
        return {
            "success": True,
            "message": f"已删除 {source}",
            "sqlite_deleted": result["sqlite_deleted"],
            "chroma_deleted": result["chroma_deleted"],
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── 统计 ───────────────────────────────────────────────────

@router.get("/stats")
async def kb_stats():
    """获取知识库统计信息"""
    try:
        from kb.database import get_kb
        kb = get_kb()
        stats = kb.get_document_stats()
        vector_stats = get_vector_stats()

        return {
            "success": True,
            "personal_db": stats,
            "vector_db": vector_stats,
            "chroma_available": is_chroma_available(),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── 检索 ───────────────────────────────────────────────────

@router.post("/search")
async def kb_search(request: Request):
    """
    混合检索知识库。

    请求体：
    {
        "query": "深度学习相关论文",
        "doc_type": "",     // 可选过滤
        "top_k": 5
    }
    """
    try:
        body = await request.json()
        query = body.get("query", "")
        doc_type = body.get("doc_type", "")
        top_k = body.get("top_k", 5)

        if not query:
            return {"success": False, "error": "缺少 query"}

        doc_type_filter = doc_type if doc_type else None
        results = hybrid_search(query, top_k=top_k, doc_type=doc_type_filter)

        # 截断过长的内容
        for r in results:
            if len(r.get("content", "")) > 2000:
                r["content"] = r["content"][:2000] + "..."

        return {
            "success": True,
            "results": results,
            "total": len(results),
            "query": query,
            "chroma_available": is_chroma_available(),
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── CV 智能解析 ───────────────────────────────────────────

@router.post("/import/parse_cv")
async def kb_parse_cv(request: Request):
    """
    用 LLM 从 CV 文本中提取结构化数据，自动导入到知识库。

    支持两种模式：
    1. 直接传 content（纯文本）
    2. 传 file_base64（PDF/DOCX），后端自动提取文本后解析

    请求体：
    {
        "content": "CV 全文文本...",       // 模式1
        "file_base64": "JVBERi0xLjQK...",  // 模式2（可选）
        "source": "cv.pdf",                // 模式2 需要，用于判断文件类型
        "llm_config": {
            "model": "gpt-4",
            "api_key": "...",
            "base_url": "...",
            "temperature": 0.1
        }
    }

    返回：
    {
        "success": true,
        "author": {"name": "Sheng Wang", "affiliation": "...", ...},
        "projects_count": 5,
        "patents_count": 2,
        "message": "已导入作者信息 + 5 个项目 + 2 条专利"
    }
    """
    try:
        body = await request.json()
        content = body.get("content", "").strip()
        file_base64 = body.get("file_base64", "")
        source = body.get("source", "")
        llm_config = body.get("llm_config", {})

        # 如果没有直接传 content，尝试从 file_base64 提取
        if not content and file_base64 and source:
            ext = os.path.splitext(source)[1].lower()
            content = _extract_text_from_file(file_base64, ext) or ""
            if content:
                print(f"[KB] parse_cv: 从文件提取文本 {len(content)} 字符")

        if not content:
            return {"success": False, "error": "缺少 content（CV 文本内容）"}

        # ── Step 1: 创建 LLM 实例 ──
        llm = _create_llm_for_parse(llm_config)

        # ── Step 2: 调用 LLM 提取结构化数据 ──
        extracted = await _extract_structured_from_cv(llm, content)

        if not extracted.get("success"):
            return {"success": False, "error": extracted.get("error", "LLM 提取失败")}

        data = extracted["data"]

        # ── Step 3: 导入到知识库 ──
        results = {}

        # 导入作者信息（先清洗姓名中的重复字符）
        author_info = data.get("author", {})
        if author_info:
            author_info["name"] = _clean_duplicate_chars(author_info.get("name", ""))
            count = import_author_info(author_info)
            results["author"] = {"imported": count, "data": author_info}

        # 导入项目列表
        projects = data.get("projects", [])
        if projects:
            count = import_projects(projects)
            results["projects"] = {"imported": count, "data": projects}

        # 导入专利列表
        patents = data.get("patents", [])
        if patents:
            count = import_patents(patents)
            results["patents"] = {"imported": count, "data": patents}

        # 导入数据集列表
        datasets = data.get("datasets", [])
        if datasets:
            count = import_datasets(datasets)
            results["datasets"] = {"imported": count, "data": datasets}

        parts = []
        if results.get("author"):
            parts.append(f"作者信息({author_info.get('name', 'N/A')})")
        if results.get("projects"):
            parts.append(f"{results['projects']['imported']} 个项目")
        if results.get("patents"):
            parts.append(f"{results['patents']['imported']} 条专利")
        if results.get("datasets"):
            parts.append(f"{results['datasets']['imported']} 个数据集")

        message = f"✅ 已导入: {', '.join(parts)}" if parts else "⚠️ CV 中未提取到可导入的结构化数据"

        print(f"[KB] CV 解析导入完成: {message}")

        return {
            "success": True,
            "message": message,
            "results": results,
            "projects_count": len(projects),
            "patents_count": len(patents),
            "datasets_count": len(datasets),
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


def _create_llm_for_parse(llm_config: Dict[str, Any]):
    """复用 agent_routes 的 LLM 创建逻辑（避免循环导入）"""
    from api.agent_routes import _create_llm
    return _create_llm(llm_config)


async def _extract_structured_from_cv(llm, content: str) -> Dict[str, Any]:
    """
    调用 LLM 从 CV 文本中提取结构化数据。

    Returns:
        {"success": True, "data": {"author": {...}, "projects": [...], ...}}
        或 {"success": False, "error": "..."}
    """
    from langchain_core.messages import HumanMessage, SystemMessage

    # 截取 CV 内容（避免 token 超限，最多 30000 字符）
    # 注意：12000 太短，会丢失后半部分的项目/专利章节（常见 CV 约 15000-20000 字符）
    cv_text = content[:30000]
    if len(content) > 30000:
        cv_text += "\n\n[... CV 内容过长，已截取前 30000 字符 ...]"

    system_prompt = """你是一个学术 CV 解析器。你的任务是从学者的 CV/简历文本中提取结构化数据。

## 严格规则
1. **只提取文本中明确出现的信息**，绝不编造任何数据
2. 如果某个字段在文本中找不到，设为空字符串 ""
3. **保持原文档的主要语言**：如果 CV 主要是中文，则 name/biography/keywords 等自由文本字段用中文；如果 CV 主要是英文则用英文
4. 列表类字段（projects、patents、datasets）如果文本中没有相关内容，返回空数组 []
5. 输出必须是合法的 JSON，不要加 markdown 代码块标记
6. **机构名称保留原文语言**（不强制翻译），但 biography/keywords 等描述性字段跟随 CV 主体语言

## 输出格式

请输出一个 JSON 对象，格式如下：

```json
{
  "author": {
    "name": "学者全名（如出现在文本中）",
    "affiliation": "当前机构/单位全称",
    "country": "国家（如明确出现）",
    "orcid": "ORCID（如有）",
    "keywords": "研究关键词，用逗号分隔",
    "biography": "一段简短的个人简介（从文本中提取，不超过 500 字）",
    "website": "个人主页 URL（如有）",
    "customFields": []
  },
  "projects": [
    {
      "title": "项目名称",
      "year": "年份",
      "funder": "资助方/来源",
      "grantId": "项目编号",
      "url": "项目链接（如有）",
      "type": "项目类型（如 NSFC, 国家自然科学基金 等）"
    }
  ],
  "patents": [
    {
      "title": "专利名称",
      "year": "年份",
      "office": "专利局",
      "number": "专利号",
      "url": "链接（如有）",
      "type": "专利类型（发明/实用新型等）"
    }
  ],
  "datasets": [
    {
      "title": "数据集名称",
      "year": "年份",
      "repo": "仓库（如 GitHub）",
      "doi": "DOI（如有）",
      "url": "链接（如有）",
      "type": "数据集类型"
    }
  ]
}
```

## 提取要点

### author
- **name（重要）**: CV 顶部的名字，注意：
  - 如果同时出现中英文版本，优先取中文姓名（不含 "Professor" 等头衔）
  - 如果名字包含重复字符（如 "王 王 王胜 胜 胜"），自动去重只保留唯一字符
  - 去掉头衔前缀（Professor, Dr., Prof. 等），只保留纯姓名
  - 示例："Professor Sheng Wang (王 王 王胜 胜 胜)" → 提取为 "Sheng Wang (王胜)" 或 "王胜"
- affiliation: 当前工作单位（通常在名字下方或 Contact 部分）
- biography: 从 Research Interests / About Me / 个人简介 部分提取
- keywords: 从 Research Interests / 研究方向 部分提取

### projects
- **重点**：必须从以下所有可能的部分提取（CV 可能有其中任意一种或多种）：
  - "Research Projects" / "Projects" / "Grants" / "Funded Research"
  - "Research Funding" / "Funding" / "RESEARCH FUNDING" （⚠️ 这个很常见！）
  - "科研项目" / "基金" / "资助" / "课题" / "研究经费"
- **每一条带金额/经费的条目都要提取为一个 project**
- 尽量提取完整的项目信息（名称、来源、时间、编号、金额）

### patents
- 从 "Patents" / "专利" 部分提取
- 只提取明确标注为专利的条目

### datasets
- 从 "Datasets" / "数据集" 部分提取
- 只提取明确标注为数据集/开源项目的条目"""

    user_msg = f"请解析以下 CV 文本，提取结构化数据。只输出 JSON，不要其他内容：\n\n---\n{cv_text}\n---"

    try:
        response = await llm.ainvoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_msg),
        ])

        raw = response.content.strip()

        # 清理可能的 markdown 代码块标记
        if raw.startswith("```"):
            # 去掉首尾的 ```json / ``` 标记
            lines = raw.split("\n")
            cleaned_lines = []
            started = False
            for line in lines:
                if line.strip().startswith("```"):
                    if started:
                        break  # 结束标记
                    started = True
                    continue
                if started:
                    cleaned_lines.append(line)
            raw = "\n".join(cleaned_lines).strip()

        import re
        # 尝试提取 JSON 块（如果 LLM 输出中混入了文字）
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if json_match:
            raw = json_match.group(0)

        data = json.loads(raw)

        # 基本校验
        if not isinstance(data, dict):
            return {"success": False, "error": f"LLM 返回的不是 JSON 对象: {type(data).__name__}"}

        # 确保关键字段存在
        data.setdefault("author", {})
        data.setdefault("projects", [])
        data.setdefault("patents", [])
        data.setdefault("datasets", [])

        if not isinstance(data["projects"], list):
            data["projects"] = []
        if not isinstance(data["patents"], list):
            data["patents"] = []
        if not isinstance(data["datasets"], list):
            data["datasets"] = []

        print(f"[KB] LLM 提取结果: author={data['author'].get('name', 'N/A')}, "
              f"projects={len(data['projects'])}, patents={len(data['patents'])}, "
              f"datasets={len(data['datasets'])}")

        return {"success": True, "data": data}

    except json.JSONDecodeError as e:
        print(f"[KB] LLM 返回的 JSON 解析失败: {e}")
        print(f"[KB] LLM 原始输出前 500 字符: {raw[:500] if 'raw' in dir() else 'N/A'}")
        return {"success": False, "error": f"LLM 返回的不是有效 JSON: {e}"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": str(e)}


# ─── 工具函数 ───────────────────────────────────────────────

def _clean_duplicate_chars(text: str) -> str:
    """清洗文本中的重复字符模式。
    
    处理 DBLP 等 API 返回的重复字问题，例如：
    "王 王 王胜 胜 胜" → "王胜"
    "Sheng Wang (王 王 王胜 胜 胜)" → "Sheng Wang (王胜)"
    
    算法：检测括号内包含的中文字符中，去除是其他字段子串的冗余字段。
    """
    if not text:
        return text
    
    import re as _re
    
    def clean_parens_content(m):
        content = m.group(1)
        # 按空格分割
        parts = content.split()
        if len(parts) <= 1:
            return m.group(0)  # 不需要清洗
        
        # 去重策略：如果一个 part 是另一个 part 的子串（且较短），则去掉较短的
        unique_parts = []
        for p in parts:
            is_redundant = False
            for other in parts:
                if p is other:
                    continue
                # 如果 p 是 other 的子串且 p 更短 → p 可能是重复的前缀
                if p in other and len(p) < len(other):
                    is_redundant = True
                    break
            if not is_redundant:
                unique_parts.append(p)
        
        return '(' + ' '.join(unique_parts) + ')'
    
    cleaned = _re.sub(r'\(([^)]+)\)', clean_parens_content, text)
    
    # 处理括号外可能存在的连续重复中文单字
    cleaned = _re.sub(r'([\u4e00-\u9fff])\s+\1\s+', r'\1 ', cleaned)
    cleaned = _re.sub(r'([\u4e00-\u9fff])\s+\1$', r'\1', cleaned)
    
    return cleaned.strip()


def _extract_text_from_file(file_base64: str, ext: str) -> str:
    """
    从 Base64 编码的文件中提取文本内容。

    支持：
    - .pdf → pypdf 提取
    - .docx → python-docx 提取
    - .txt/.md/.csv 等纯文本 → 直接解码 UTF-8

    Returns:
        提取的纯文本内容
    """
    try:
        raw_bytes = base64.b64decode(file_base64)
    except Exception:
        return ""

    ext = ext.lstrip(".")

    if ext == "pdf":
        return _extract_pdf_text(raw_bytes)
    elif ext == "docx":
        return _extract_docx_text(raw_bytes)
    else:
        # 纯文本文件，尝试 UTF-8 解码
        for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                return raw_bytes.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return ""


def _extract_pdf_text(raw_bytes: bytes) -> str:
    """使用 pypdf 从 PDF 二进制数据中提取文本"""
    try:
        from io import BytesIO
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw_bytes))
        texts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                texts.append(page_text.strip())

        full_text = "\n\n".join(texts)
        if full_text.strip():
            print(f"[KB] PDF 提取成功: {len(reader.pages)} 页, {len(full_text)} 字符")
        else:
            print(f"[KB] PDF 提取警告: {len(reader.pages)} 页但未提取到文本（可能是扫描件）")
        return full_text
    except ImportError:
        print("[KB] pypdf 未安装，无法提取 PDF 文本")
        return ""
    except Exception as e:
        print(f"[KB] PDF 文本提取失败: {e}")
        return ""


def _extract_docx_text(raw_bytes: bytes) -> str:
    """使用 python-docx 从 DOCX 二进制数据中提取文本"""
    try:
        from io import BytesIO
        from docx import Document

        doc = Document(BytesIO(raw_bytes))
        texts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                texts.append(para.text.strip())

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)

        full_text = "\n".join(texts)
        print(f"[KB] DOCX 提取成功: {len(doc.paragraphs)} 段, {len(full_text)} 字符")
        return full_text
    except ImportError:
        print("[KB] python-docx 未安装，无法提取 DOCX 文本")
        return ""
    except Exception as e:
        print(f"[KB] DOCX 文本提取失败: {e}")
        return ""


def _chunk_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> List[str]:
    """
    简单文本分块（按字符数，尽量在段落边界分割）。
    """
    if not text or not text.strip():
        return []

    # 先按段落分割
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current_chunk) + len(para) + 1 <= chunk_size:
            current_chunk = (current_chunk + "\n" + para).strip()
        else:
            if current_chunk:
                chunks.append(current_chunk)

            # 如果单段超过 chunk_size，强制分割
            if len(para) > chunk_size:
                for i in range(0, len(para), chunk_size - chunk_overlap):
                    chunk = para[i:i + chunk_size]
                    if chunk.strip():
                        chunks.append(chunk.strip())
                current_chunk = ""
            else:
                current_chunk = para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks
